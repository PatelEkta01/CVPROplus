import pdfplumber
import docx
import re
import spacy
import google.generativeai as genai
import phonenumbers
import json
import os
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

gemini_api_key = os.getenv("REACT_APP_GEMINI_API_KEY")

# Configure Gemini API with hardcoded key
genai.configure(api_key=gemini_api_key)

def extract_text_from_pdf(pdf_file):
    """
    Extract text from a PDF file using pdfplumber
    """
    try:
        text = ""
        with pdfplumber.open(pdf_file) as pdf:
            for i, page in enumerate(pdf.pages):
                print(f"Processing page {i+1} of {len(pdf.pages)}")
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                else:
                    print(f"Warning: No text extracted from page {i+1}")
        
        if not text.strip():
            print("Warning: No text extracted from the entire PDF")
            
        return text.strip()
    except Exception as e:
        print(f"Error extracting text from PDF: {str(e)}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(docx_file):
    doc = docx.Document(docx_file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text.strip()

def parse_resume_text(text):
    """Basic parsing to extract resume fields (can be improved with NLP)"""
    resume_data = {
        "name": None,
        "email": None,
        "phone": None,
        "summary": None,
        "Experience": [],
        "skills": [],
        "projects": []
    }

    lines = text.split("\n")
    for line in lines:
        if "@" in line:  # Extract Email
            resume_data["email"] = line.strip()
        elif any(char.isdigit() for char in line) and len(line) >= 10:  # Extract Phone Number
            resume_data["phone"] = line.strip()
        elif "Experience" in line:
            resume_data["Experience"].append(line.strip())
        elif "Skills" in line:
            resume_data["skills"].append(line.strip())
        elif "Projects" in line:
            resume_data["projects"].append(line.strip())
        elif resume_data["name"] is None:  # Assume the first line is the name
            resume_data["name"] = line.strip()
    
    return resume_data

nlp = spacy.load("en_core_web_sm")  # Load NLP model for text processing

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file"""
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file"""
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_email(text):
    """Extract email from text using regex"""
    match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else None

def extract_phone(text):
    """Extract phone number from text using phonenumbers library"""
    matches = re.findall(r'\+?\d[\d -]{8,15}\d', text)
    for match in matches:
        try:
            phone = phonenumbers.parse(match, None)
            if phonenumbers.is_valid_number(phone):
                return phonenumbers.format_number(phone, phonenumbers.PhoneNumberFormat.INTERNATIONAL)
        except:
            continue
    return None

def extract_name(text):
    """Extract the most probable name using NLP"""
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            return ent.text
    return None

def extract_sections(text):
    """Extract structured sections from resume"""
    sections = {
        "experience": [],
        "skills": [],
        "projects": [],
    }
    
    lines = text.split("\n")
    current_section = None

    for line in lines:
        line_lower = line.lower().strip()

        if "experience" in line_lower:
            current_section = "experience"
        elif "skills" in line_lower:
            current_section = "skills"
        elif "projects" in line_lower:
            current_section = "projects"
        elif current_section:
            sections[current_section].append(line.strip())

    return sections

def parse_resume(text):
    return {
        "name": extract_name(text),
        "email": extract_email(text),
        "phone": extract_phone(text),
        **extract_sections(text),
    }

def normalize_spaces(text):
    """Normalize spaces in the text to ensure proper formatting."""
    return " ".join(text.split())

def parse_resume_with_gemini(text):
    """Uses Gemini AI to extract structured data from the resume."""
    model = genai.GenerativeModel("gemini-2.0-flash")  # Ensure correct model usage
    chat = model.start_chat()  # Start a conversation

    normalized_text = normalize_spaces(text)

    prompt = f"""
    Extract the following details from this resume:
    - Name
    - Email
    - Phone Number only in 10 digits
    - Address
    - Summary (Ensure it's extracted properly. If missing, return null.)
    - Skills (as a list)
    - Experience (Job Title, Company, Start Date, End Date, Location, Description as separate tasks)
    - Projects (Title, Description as separate tasks, Technologies used)
    - Education (Institution, Graduation Date, Course, Location)

    Structure it in **valid JSON** format:
    {json.dumps({
        "personal":{
                    "name": "string",
                    "email": "string",
                    "phone": "string",
                    "address": "string",
                    "summary": "string or null",
                    },
        "skills": ["string"],
        "experience": [
            {
                "jobTitle": "string",
                "company": "string",
                "startDate": "MM/YYYY",
                "endDate": "MM/YYYY or null (if current)",
                "location": "string",
                "tasks": ["string"]  
            }
        ],
        "projects": [
            {
                "name": "string",
                "tasks": ["string"],  
                "technologies": ["string"]
            }
        ],
        "education": [
            {
                "institution": "string",
                "graduation_date": "MM/YYYY",
                "course": "string",
                "location": "string"
            }
        ]
    }, indent=5)}

    Resume Text:
    {normalized_text}
    """

    response = chat.send_message(prompt)
    response_text = response.text.strip()

    if response_text.startswith("```json") and response_text.endswith("```"):
        response_text = response_text[7:-3].strip()  

    try:
        response_dict = json.loads(response_text)
        return response_dict
    except json.JSONDecodeError as e:
        print(f"Error parsing JSON response: {e}")
        return None
